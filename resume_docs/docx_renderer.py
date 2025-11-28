"""DOCX rendering helpers."""
from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable

from . import constants, models

SECTION_LABELS = {
    "zh-CN": {
        "skills": "技能",
        "projects": "项目经验",
        "work": "工作经历",
        "roles": "角色亮点",
        "education": "教育背景",
        "certifications": "认证",
    },
    "en-US": {
        "skills": "Skills",
        "projects": "Projects",
        "work": "Experience",
        "roles": "Role Highlights",
        "education": "Education",
        "certifications": "Certifications",
    },
}


class MissingDependencyError(RuntimeError):
    """Raised when an optional dependency is unavailable."""


def _ensure_docx_imports():
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError as exc:  # pragma: no cover - runtime dependency
        raise MissingDependencyError(
            "python-docx is required for DOCX rendering. Install requirements.txt deps"
        ) from exc
    return Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor


def render_docx(
    resume: models.ResumeDocument,
    output_path: Path,
    theme_name: str,
    locale: str,
    include_contact: bool = False,
) -> Path:
    Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor = _ensure_docx_imports()
    theme = constants.THEMES[theme_name]
    labels = SECTION_LABELS.get(locale, SECTION_LABELS[constants.DEFAULT_LOCALE])
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    _render_header(doc, resume.personal_info, include_contact, theme, WD_ALIGN_PARAGRAPH, Pt, RGBColor)
    _render_skills(doc, resume.skills, labels["skills"], theme, Pt, RGBColor)
    _render_projects(doc, resume.projects, labels["projects"], theme, Pt, RGBColor)
    _render_work(doc, resume.work, labels, theme, Pt, RGBColor)
    _render_education_and_certifications(doc, resume.personal_info, labels, theme, Pt, RGBColor)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _render_header(doc, info: models.PersonalInfo, include_contact: bool, theme, WD_ALIGN_PARAGRAPH, Pt, RGBColor):
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(info.name)
    name_run.font.size = Pt(28)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name_run.font.name = "Calibri"

    # Add bottom border line
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = name_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    if include_contact:
        contact_parts = [value for value in [info.phone, info.email, info.address] if value]
    else:
        contact_parts = [value for value in [info.email] if value]
    if info.github:
        contact_parts.append(info.github)
    if contact_parts:
        for contact_item in contact_parts:
            contact_para = doc.add_paragraph(contact_item)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            contact_para.paragraph_format.space_before = Pt(3)
            contact_para.paragraph_format.space_after = Pt(3)
            for run in contact_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "Calibri"


def _render_section_heading(doc, text: str, theme, Pt, RGBColor):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Calibri"
    return heading


def _render_skills(doc, skills: models.SkillsSummary, section_label: str, theme, Pt, RGBColor):
    if not skills.categories:
        return
    _render_section_heading(doc, section_label, theme, Pt, RGBColor)
    for category in skills.categories:
        para = doc.add_paragraph()
        para.style = doc.styles["Normal"]
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(f"{category.category}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Calibri"
        if category.items:
            items_run = para.add_run(" · ".join(category.items))
            items_run.font.name = "Calibri"


def _render_projects(doc, projects: Iterable[models.Project], section_label: str, theme, Pt, RGBColor):
    if not projects:
        return
    _render_section_heading(doc, section_label, theme, Pt, RGBColor)
    for project in projects:
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(2)
        title_run = title_para.add_run(project.project_name)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_run.font.name = "Calibri"
        if project.company_or_context:
            ctx_run = title_para.add_run(f" | {project.company_or_context}")
            ctx_run.font.name = "Calibri"
        if project.timeframe and project.timeframe.label:
            time_run = title_para.add_run(f" ({project.timeframe.label})")
            time_run.font.name = "Calibri"
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_after = Pt(3)
        meta_items = []
        if project.role_title:
            meta_items.append(project.role_title)
        if project.role_perspective:
            meta_items.append(project.role_perspective)
        if project.data_domain:
            meta_items.append(project.data_domain)
        if meta_items:
            meta_run = meta_para.add_run(" | ".join(meta_items))
            meta_run.font.name = "Calibri"
        if project.project_overview:
            overview_para = doc.add_paragraph(project.project_overview)
            overview_para.style = doc.styles["Normal"]
            overview_para.paragraph_format.space_after = Pt(3)
            for run in overview_para.runs:
                run.font.name = "Calibri"
        for metrics_label, values in project.impact_metrics.grouped():
            if values:
                metric_para = doc.add_paragraph()
                metric_para.paragraph_format.space_after = Pt(2)
                run = metric_para.add_run(f"{metrics_label}: ")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "Calibri"
                values_run = metric_para.add_run("; ".join(values))
                values_run.font.name = "Calibri"
        bullet_fields = [
            ("Challenges", project.challenges_or_objectives),
            ("Responsibilities", project.responsibilities),
            ("Solution", project.architecture_or_solution),
            ("Deliverables", project.deliverables_or_features),
            ("Impact", project.metrics_or_impact),
        ]
        for label, values in bullet_fields:
            if not values:
                continue
            label_para = doc.add_paragraph(f"{label}:")
            label_para.style = doc.styles["Normal"]
            label_para.paragraph_format.space_before = Pt(3)
            label_para.paragraph_format.space_after = Pt(2)
            for run in label_para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "Calibri"
            for value in values:
                bullet_para = doc.add_paragraph(value, style="List Bullet")
                for run in bullet_para.runs:
                    run.font.name = "Calibri"


def _render_work(doc, work: models.WorkSummary, labels: Dict[str, str], theme, Pt, RGBColor):
    if work.experiences:
        _render_section_heading(doc, labels["work"], theme, Pt, RGBColor)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.autofit = False
        table.allow_autofit = False
        header_cells = table.rows[0].cells
        header_cells[0].text = "Company"
        header_cells[1].text = "Duration"
        header_cells[2].text = "Title"
        for cell in header_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = "Calibri"
        for exp in work.experiences:
            row_cells = table.add_row().cells
            row_cells[0].text = exp.company or ""
            row_cells[1].text = exp.duration or ""
            row_cells[2].text = exp.title or ""
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Calibri"


def _render_education_and_certifications(doc, info: models.PersonalInfo, labels: Dict[str, str], theme, Pt, RGBColor):
    if info.education:
        _render_section_heading(doc, labels["education"], theme, Pt, RGBColor)
        for entry in info.education:
            edu_para = doc.add_paragraph(entry.description, style="List Bullet")
            for run in edu_para.runs:
                run.font.name = "Calibri"
    if info.certifications:
        _render_section_heading(doc, labels["certifications"], theme, Pt, RGBColor)
        for cert in info.certifications:
            para = doc.add_paragraph(style="List Bullet")
            name_run = para.add_run(cert.name)
            name_run.font.name = "Calibri"
            if cert.url:
                url_run = para.add_run(f" ({cert.url})")
                url_run.font.name = "Calibri"
